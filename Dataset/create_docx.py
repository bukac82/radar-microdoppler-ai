from docx import Document
from docx.shared import Pt, Inches

def create_readme():
    doc = Document()
    
    # Title
    doc.add_heading('Helicopter Micro-Doppler Dataset README', 0)
    
    # Intro
    doc.add_paragraph(
        "This dataset collection contains synthetic micro-Doppler signatures for helicopters, "
        "generated based on the mathematical models established by Dr. Victor C. Chen. "
        "The data is intended for training machine learning models to classify helicopter types "
        "and extract micro-motion features."
    )
    
    # Section 1
    doc.add_heading('1. Base Dataset', level=1)
    doc.add_paragraph("File: helicopter_microdoppler_dataset.csv\n"
                      "Total Samples: 100,000 rows\n"
                      "Columns per Sample: 1,005\n"
                      "File Size: ~2.1 GB")
    
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("Radar frequency (10 GHz), elevation (10 deg), and bulk target speed (0 m/s) are held constant in this dataset to isolate the micro-Doppler signatures.")
    
    # Section 2
    doc.add_heading('2. Extended Dataset', level=1)
    doc.add_paragraph("File: helicopter_microdoppler_extended_dataset.csv\n"
                      "Total Samples: 100,000 rows\n"
                      "Columns per Sample: 1,008\n"
                      "File Size: ~2.1 GB")
    
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("This dataset includes randomized values for radar frequency, elevation, and bulk target motion (Doppler shift) to provide a more robust training distribution.")
    
    # Section 3
    doc.add_heading('3. Data Structure (Extended Dataset)', level=1)
    doc.add_paragraph("Every row in the CSV represents a single 0.5-second simulation (sampled at 1000 Hz) with the following columns:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column Name'
    hdr_cells[1].text = 'Description'
    
    rows = [
        ('num_blades', 'Number of rotor blades (2, 3, or 4).'),
        ('radius_m', 'The length/radius of the blade in meters.'),
        ('rpm', 'The rotations per minute of the rotor.'),
        ('tip_velocity_m_s', 'The calculated tip velocity in meters per second.'),
        ('snr_db', 'The Signal-to-Noise Ratio (dB) applied to the sample.'),
        ('radar_freq_ghz', '[NEW] The radar carrier frequency (8.0 to 12.0 GHz).'),
        ('radar_elevation_deg', '[NEW] The elevation angle of the radar (0 to 45 degrees).'),
        ('target_speed_m_s', '[NEW] Target radial speed (-50 to 50 m/s), introducing bulk Doppler shift.'),
        ('I_0 to I_499', 'The In-Phase (Real) component of the 500-sample time series.'),
        ('Q_0 to Q_499', 'The Quadrature (Imaginary) component of the 500-sample time series.')
    ]
    
    for name, desc in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc

    # Section 4
    doc.add_heading('4. Helicopter Parameters', level=1)
    doc.add_paragraph(
        "Parameters were randomly sampled within realistic ranges to ensure dataset variance:\n\n"
        "• 2-Blade (e.g., Bell UH-1): RPM [300, 350], Radius [6.5m, 7.5m]\n"
        "• 3-Blade (e.g., Gazelle): RPM [360, 400], Radius [4.5m, 5.5m]\n"
        "• 4-Blade (e.g., AH-64 Apache): RPM [250, 300], Radius [7.0m, 8.5m]\n\n"
        "Tip velocity is dynamically calculated as V_tip = RPM * (2π/60) * Radius."
    )
    
    doc.save('README_Helicopter_MicroDoppler.docx')
    print("DOCX created successfully.")

if __name__ == '__main__':
    create_readme()
